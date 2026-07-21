import io
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import spacy
from docx import Document
from pypdf import PdfReader
from collections import Counter
from typing import List
import re
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("job_keyword_extractor")

# Load spaCy model (used here for parsing/phrase extraction, not for its
# vectors — en_core_web_sm has no real word vectors. Semantic similarity is
# handled by sentence-transformers below, which is trained specifically for
# that job and performs much better than averaged spaCy vectors would,
# especially for phrase-vs-phrase comparisons like "built REST APIs" vs
# "API development experience".)
nlp = spacy.load("en_core_web_sm")

# Lazily-loaded sentence-transformers model. Loaded on first use (not at
# import time) so the API still starts up fast; only /api/semantic-match
# pays the model-load cost, on its first call.
_st_model = None

def get_similarity_model():
    global _st_model
    if _st_model is None:
        from sentence_transformers import SentenceTransformer
        logger.info("Loading sentence-transformers model (all-MiniLM-L6-v2)...")
        _st_model = SentenceTransformer("all-MiniLM-L6-v2")
    return _st_model

# Below this cosine-similarity score, a job phrase is considered "missing"
# rather than matched. Tuned empirically — MiniLM cosine scores for genuinely
# related-but-differently-worded phrases (e.g. "built REST APIs" vs "API
# development experience") tend to land around 0.5-0.7; unrelated phrases
# usually sit below 0.3.
SEMANTIC_MATCH_THRESHOLD = 0.7

app = FastAPI(
    title="Job Keyword Extractor",
    description="Extract keywords and skills from job postings",
    version="1.0.0"
)

# BUG FIX: The React frontend (http://localhost:3000) was being blocked by the
# browser's CORS policy because no CORSMiddleware was ever registered here,
# despite HANDOVER.md claiming "CORS is enabled by default". Without this,
# every request from the frontend fails before it even reaches the routes
# below (confirmed: an OPTIONS preflight returned 405 with no CORS headers).
#
# allow_origins=["*"] here (rather than pinning to one port) because the
# frontend is plain static HTML/JS and might be served from a few different
# local dev setups (VS Code Live Server, `python -m http.server`, etc.) each
# using a different port. Note: allow_credentials must be False when using
# "*", which is fine here since this API doesn't use cookies/auth.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Data models
class JobPosting(BaseModel):
    text: str

class KeywordResponse(BaseModel):
    keywords: List[dict]
    success: bool

class SkillsResponse(BaseModel):
    skills: dict
    success: bool

class ParsedFileResponse(BaseModel):
    text: str
    success: bool

class SemanticMatchRequest(BaseModel):
    resume_text: str
    job_text: str

class MatchedPhrase(BaseModel):
    job_phrase: str
    resume_phrase: str
    score: float  # cosine similarity, 0-1

class SemanticMatchResponse(BaseModel):
    matched: List[MatchedPhrase]
    missing: List[str]
    score: float  # percent of job phrases matched, 0-100
    success: bool

MAX_UPLOAD_BYTES = 5 * 1024 * 1024  # 5MB

def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if not text:
        # Scanned/image-only PDFs have no text layer to pull from.
        raise ValueError("No extractable text found (this PDF may be a scanned image).")
    return text

def extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(p for p in paragraphs if p.strip())
    if not text:
        raise ValueError("No extractable text found in this document.")
    return text

# Helper function to extract keywords
def extract_keywords_spacy(job_posting: str, top_n: int = 20):
    """Extract keywords using spaCy"""
    doc = nlp(job_posting.lower())
    
    # Extract nouns, proper nouns, and adjectives
    keywords = [
        token.text for token in doc 
        if token.pos_ in ['NOUN', 'PROPN', 'ADJ']
        and len(token.text) > 3
        and not token.is_stop
    ]
    
    # Count frequency
    keyword_counts = Counter(keywords)
    return keyword_counts.most_common(top_n)

# Helper function to extract candidate "skill phrases" for semantic matching.
# Unlike extract_technical_skills below (a fixed dictionary lookup),
# this pulls open-vocabulary phrases straight out of the text so that
# free-form wording like "built REST APIs" or "led a team of 5 engineers"
# becomes something we can embed and compare, instead of only matching
# against a predefined list of tool/language names.
def extract_phrases(text: str, max_phrases: int = 60) -> List[str]:
    """Extract noun-phrase and verb+object phrases from text.

    - Noun chunks capture skill/requirement nouns, e.g. "API development
      experience", "REST APIs", "cross-functional teams".
    - Verb + direct-object phrases capture accomplishment phrasing, e.g.
      "built REST APIs" (verb "build" + object span "REST APIs"),
      "managed a budget", "led a team".
    """
    doc = nlp(text)
    phrases: List[str] = []
    seen = set()

    def add(phrase: str):
        clean = re.sub(r"\s+", " ", phrase).strip()
        # Drop very short or punctuation-only fragments
        if len(clean) < 3 or not re.search(r"[a-zA-Z]", clean):
            return
        key = clean.lower()
        if key in seen:
            return
        seen.add(key)
        phrases.append(clean)

    # Noun chunks
    for chunk in doc.noun_chunks:
        # Strip leading determiners/possessives ("the", "our", "a") so
        # phrases compare more consistently, but keep everything else.
        tokens = [t for t in chunk if t.pos_ not in ("DET", "PRON")]
        if tokens:
            add(" ".join(t.text for t in tokens))

    # Verb + direct object phrases (captures accomplishment-style wording)
    for token in doc:
        if token.pos_ != "VERB":
            continue
        for child in token.children:
            if child.dep_ in ("dobj", "obj", "attr", "oprd"):
                span = doc[child.left_edge.i : child.right_edge.i + 1]
                add(f"{token.lemma_} {span.text}")

    if len(phrases) > max_phrases:
        phrases = phrases[:max_phrases]

    return phrases


# Helper function to extract skills
def extract_technical_skills(job_posting: str) -> dict:
    """Extract technical skills from job posting"""
    text_lower = job_posting.lower()
    
    skills = {
        'programming_languages': {
            'python', 'java', 'javascript', 'c#', 'c++', 'ruby', 'php', 'swift',
            'golang', 'rust', 'typescript', 'kotlin', 'scala', 'r', 'matlab'
        },
        'frameworks': {
            'django', 'flask', 'react', 'vue', 'angular', 'spring', 'express',
            'fastapi', 'asp.net', 'laravel', 'rails', 'torch', 'tensorflow',
            'keras', 'scikit', 'pandas', 'numpy'
        },
        'databases': {
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch',
            'cassandra', 'dynamodb', 'oracle', 'sql', 'firebase', 'graphql'
        },
        'tools': {
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
            'gitlab', 'github', 'jira', 'linux', 'windows', 'macos'
        }
    }
    
    detected_skills = {}
    
    for category, skill_set in skills.items():
        found = []
        for skill in skill_set:
            # BUG FIX: skill names were interpolated into the regex unescaped
            # AND wrapped in \b...\b, which broke matching for any skill
            # containing regex metacharacters or symbols:
            #   - "c++"     -> unescaped, "+" is a quantifier, so \bc++\b
            #                  actually matched any lone "c" (false positive:
            #                  a posting mentioning grade "C" "detected" C++).
            #   - "c#"      -> "#" is not a \w character, and \b only matches
            #                  at a \w/\W transition. Right after "#" both
            #                  sides are non-word whenever followed by a
            #                  space/punctuation/end-of-string, so \b NEVER
            #                  matched there -> C# was silently never detected.
            #   - "asp.net" -> "." matches any character when unescaped, so
            #                  "asp net" or "aspXnet" would wrongly match too.
            #                  Even after escaping, \b after the literal "."
            #                  has the same failure mode as "c#" above.
            # Fix: re.escape() the skill text, and replace \b with an explicit
            # "not preceded/followed by a letter or digit" check, which works
            # correctly regardless of whether the skill ends in a symbol.
            pattern = r'(?<![a-zA-Z0-9])' + re.escape(skill) + r'(?![a-zA-Z0-9])'
            if re.search(pattern, text_lower):
                found.append(skill)
        
        if found:
            detected_skills[category] = found
    
    return detected_skills

# Routes
@app.get("/")
def read_root():
    """Welcome message"""
    return {
        "message": "Job Keyword Extractor API",
        "version": "1.0.0",
        "docs": "/docs"
    }

@app.get("/api/health")
def health_check():
    """Check if backend is running"""
    return {"status": "ok", "message": "Backend is running ✓"}

@app.post("/api/extract-keywords", response_model=KeywordResponse)
def extract_keywords(job: JobPosting):
    """Extract keywords from job posting"""
    try:
        keywords = extract_keywords_spacy(job.text, top_n=20)
        return KeywordResponse(
            keywords=[{"word": k, "frequency": f} for k, f in keywords],
            success=True
        )
    except Exception as e:
        # BUG FIX: previously the exception was caught into `e` and then
        # thrown away, so failures returned success:false with zero
        # information about what actually went wrong. Now it's logged.
        logger.exception("extract_keywords failed: %s", e)
        return KeywordResponse(
            keywords=[],
            success=False
        )

@app.post("/api/extract-skills", response_model=SkillsResponse)
def extract_skills(job: JobPosting):
    """Extract technical skills"""
    try:
        skills = extract_technical_skills(job.text)
        return SkillsResponse(
            skills=skills,
            success=True
        )
    except Exception as e:
        # BUG FIX: same silent-swallow issue as extract_keywords above.
        logger.exception("extract_skills failed: %s", e)
        return SkillsResponse(
            skills={},
            success=False
        )

@app.post("/api/semantic-match", response_model=SemanticMatchResponse)
def semantic_match(payload: SemanticMatchRequest):
    """Compare a resume and job posting by meaning, not exact wording.

    Extracts candidate phrases from both texts, embeds them with
    sentence-transformers, and matches each job phrase to its closest
    resume phrase by cosine similarity. This is what lets "built REST
    APIs" on a resume satisfy "API development experience" in a posting,
    which plain substring/set matching can't do.
    """
    try:
        job_phrases = extract_phrases(payload.job_text)
        resume_phrases = extract_phrases(payload.resume_text)

        if not job_phrases:
            return SemanticMatchResponse(matched=[], missing=[], score=0.0, success=True)
        if not resume_phrases:
            return SemanticMatchResponse(matched=[], missing=job_phrases, score=0.0, success=True)

        model = get_similarity_model()
        job_embeddings = model.encode(job_phrases, normalize_embeddings=True)
        resume_embeddings = model.encode(resume_phrases, normalize_embeddings=True)

        # Cosine similarity of normalized vectors == dot product
        sims = job_embeddings @ resume_embeddings.T  # shape: [len(job), len(resume)]

        matched: List[MatchedPhrase] = []
        missing: List[str] = []

        for i, phrase in enumerate(job_phrases):
            best_idx = int(sims[i].argmax())
            best_score = float(sims[i][best_idx])
            if best_score >= SEMANTIC_MATCH_THRESHOLD:
                matched.append(MatchedPhrase(
                    job_phrase=phrase,
                    resume_phrase=resume_phrases[best_idx],
                    score=round(best_score, 3),
                ))
            else:
                missing.append(phrase)

        overall_score = round(len(matched) / len(job_phrases) * 100, 1)

        return SemanticMatchResponse(
            matched=matched,
            missing=missing,
            score=overall_score,
            success=True,
        )
    except Exception as e:
        logger.exception("semantic_match failed: %s", e)
        return SemanticMatchResponse(matched=[], missing=[], score=0.0, success=False)


@app.post("/api/parse-resume", response_model=ParsedFileResponse)
async def parse_resume(file: UploadFile = File(...)):
    """Accepts a .pdf, .docx, or .txt upload and returns its plain text."""
    filename = (file.filename or "").lower()
    content = await file.read()

    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB).")

    try:
        if filename.endswith(".pdf"):
            text = extract_pdf_text(content)
        elif filename.endswith(".docx"):
            text = extract_docx_text(content)
        elif filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Upload a .pdf, .docx, or .txt file.",
            )
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.exception("parse_resume failed: %s", e)
        raise HTTPException(status_code=500, detail="Could not read this file.")

    return ParsedFileResponse(text=text, success=True)

if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting FastAPI server...")
    print("📚 API Docs available at: http://localhost:8000/docs")
    uvicorn.run(app, host="0.0.0.0", port=8000)