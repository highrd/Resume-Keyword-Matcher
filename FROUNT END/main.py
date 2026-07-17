# backend/main.py
import io
import logging
import re
from collections import Counter
from typing import List

import spacy
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from pypdf import PdfReader

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("job_keyword_extractor")

nlp = spacy.load("en_core_web_sm")

app = FastAPI(
    title="Job Keyword Extractor",
    description="Extract keywords and skills from job postings, and parse resume uploads",
    version="1.1.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------- Data models ----------

class JobPosting(BaseModel):
    text: str


class KeywordResponse(BaseModel):
    keywords: List[dict]
    success: bool


class SkillsResponse(BaseModel):
    skills: dict
    success: bool


class ParsedFileResponse(BaseModel):
    text: str
    success: bool


# ---------- Keyword / skill extraction (unchanged from test_fastapi.py) ----------

def extract_keywords_spacy(job_posting: str, top_n: int = 20):
    doc = nlp(job_posting.lower())
    keywords = [
        token.text
        for token in doc
        if token.pos_ in ["NOUN", "PROPN", "ADJ"]
        and len(token.text) > 3
        and not token.is_stop
    ]
    keyword_counts = Counter(keywords)
    return keyword_counts.most_common(top_n)


def extract_technical_skills(job_posting: str) -> dict:
    text_lower = job_posting.lower()

    skills = {
        "programming_languages": {
            "python", "java", "javascript", "c#", "c++", "ruby", "php", "swift",
            "golang", "rust", "typescript", "kotlin", "scala", "r", "matlab",
        },
        "frameworks": {
            "django", "flask", "react", "vue", "angular", "spring", "express",
            "fastapi", "asp.net", "laravel", "rails", "torch", "tensorflow",
            "keras", "scikit", "pandas", "numpy",
        },
        "databases": {
            "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
            "cassandra", "dynamodb", "oracle", "sql", "firebase", "graphql",
        },
        "tools": {
            "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "git",
            "gitlab", "github", "jira", "linux", "windows", "macos",
        },
    }

    detected_skills = {}
    for category, skill_set in skills.items():
        found = []
        for skill in skill_set:
            pattern = r"(?<![a-zA-Z0-9])" + re.escape(skill) + r"(?![a-zA-Z0-9])"
            if re.search(pattern, text_lower):
                found.append(skill)
        if found:
            detected_skills[category] = found

    return detected_skills


# ---------- File parsing helpers ----------

MAX_UPLOAD_BYTES = 5 * 1024 * 1024  # 5MB

def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if not text:
        # Happens with scanned/image-only PDFs — there's no text layer to pull.
        raise ValueError("No extractable text found (this PDF may be a scanned image).")
    return text


def extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs]
    # also pick up text inside tables, which resumes sometimes use for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(p for p in paragraphs if p.strip())
    if not text:
        raise ValueError("No extractable text found in this document.")
    return text


# ---------- Routes ----------

@app.get("/")
def read_root():
    return {"message": "Job Keyword Extractor API", "version": "1.1.0", "docs": "/docs"}


@app.get("/api/health")
def health_check():
    return {"status": "ok", "message": "Backend is running"}


@app.post("/api/extract-keywords", response_model=KeywordResponse)
def extract_keywords(job: JobPosting):
    try:
        keywords = extract_keywords_spacy(job.text, top_n=20)
        return KeywordResponse(
            keywords=[{"word": k, "frequency": f} for k, f in keywords], success=True
        )
    except Exception as e:
        logger.exception("extract_keywords failed: %s", e)
        return KeywordResponse(keywords=[], success=False)


@app.post("/api/extract-skills", response_model=SkillsResponse)
def extract_skills(job: JobPosting):
    try:
        skills = extract_technical_skills(job.text)
        return SkillsResponse(skills=skills, success=True)
    except Exception as e:
        logger.exception("extract_skills failed: %s", e)
        return SkillsResponse(skills={}, success=False)


@app.post("/api/parse-resume", response_model=ParsedFileResponse)
async def parse_resume(file: UploadFile = File(...)):
    """Accepts a .pdf, .docx, or .txt upload and returns its plain text."""
    filename = (file.filename or "").lower()
    content = await file.read()

    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB).")

    try:
        if filename.endswith(".pdf"):
            text = extract_pdf_text(content)
        elif filename.endswith(".docx"):
            text = extract_docx_text(content)
        elif filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Upload a .pdf, .docx, or .txt file.",
            )
    except HTTPException:
        raise
    except ValueError as e:
        # No text layer / empty document — a real problem with the file, not the server
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.exception("parse_resume failed: %s", e)
        raise HTTPException(status_code=500, detail="Could not read this file.")

    return ParsedFileResponse(text=text, success=True)


if __name__ == "__main__":
    import uvicorn

    print("🚀 Starting FastAPI server...")
    print("📚 API Docs available at: http://localhost:8000/docs")
    uvicorn.run(app, host="0.0.0.0", port=8000)
